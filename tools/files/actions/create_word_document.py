"""
Creates a Word document (.docx) from text/markdown content.
"""
import os
import logging
from typing import Any, Optional
from pydantic import BaseModel, Field
from pathlib import Path
from docx import Document
from tools.base import BaseTool
from tools.manifest import ToolManifest
from tools.web.actions.download_media import resolve_directory

logger = logging.getLogger("CreateWordDocumentTool")

class CreateWordDocumentInput(BaseModel):
    title: str = Field(..., description="The title of the Word document.")
    content: str = Field(..., description="The content of the document (supports markdown-style headings #, ##, ###).")
    output_dir: Optional[str] = Field(None, description="The folder to save the document in (e.g., 'documentos', 'descargas', or a full path).")
    filename: str = Field(..., description="The filename for the Word document (e.g., 'reporte_noticias.docx').")

class CreateWordDocumentTool(BaseTool):
    manifest = ToolManifest(
        name="create_word_document",
        description="Creates a Microsoft Word document (.docx) with structured headings and paragraphs from text/markdown content.",
        arguments_schema=CreateWordDocumentInput,
        permission_level="low",
        risk="modification"
    )

    async def execute(self, **kwargs) -> Any:
        title = kwargs.get("title")
        content = kwargs.get("content")
        output_dir_str = kwargs.get("output_dir")
        filename = kwargs.get("filename")
        
        if not title or not content or not filename:
            return {"error": "Missing parameters: title, content, and filename are required."}
            
        # Ensure filename ends with .docx
        if not filename.endswith(".docx"):
            filename += ".docx"
            
        output_dir = resolve_directory(output_dir_str, default_type="downloads")
        file_path = output_dir / filename
        
        try:
            logger.info(f"Generating Word document: {file_path}")
            doc = Document()
            
            # Document Title
            doc.add_heading(title, level=0)
            
            # Parse markdown-like content line by line
            lines = content.split("\n")
            current_paragraph_lines = []
            
            def flush_paragraph():
                nonlocal current_paragraph_lines
                if current_paragraph_lines:
                    p_text = " ".join(current_paragraph_lines).strip()
                    if p_text:
                        doc.add_paragraph(p_text)
                    current_paragraph_lines = []
            
            for line in lines:
                line_stripped = line.strip()
                if not line_stripped:
                    flush_paragraph()
                    continue
                    
                # Headings
                if line_stripped.startswith("# "):
                    flush_paragraph()
                    doc.add_heading(line_stripped[2:], level=1)
                elif line_stripped.startswith("## "):
                    flush_paragraph()
                    doc.add_heading(line_stripped[3:], level=2)
                elif line_stripped.startswith("### "):
                    flush_paragraph()
                    doc.add_heading(line_stripped[4:], level=3)
                else:
                    # Accumulate paragraph text
                    current_paragraph_lines.append(line_stripped)
                    
            # Flush any remaining paragraph
            flush_paragraph()
            
            # Save the file
            doc.save(str(file_path))
            
            return {
                "status": "success",
                "file_path": str(file_path),
                "filename": filename,
                "output_dir": str(output_dir),
                "message": f"Documento Word creado correctamente en: {file_path}"
            }
            
        except Exception as e:
            logger.error(f"Failed to create Word document: {e}")
            return {"error": f"Failed to create Word document: {str(e)}"}
